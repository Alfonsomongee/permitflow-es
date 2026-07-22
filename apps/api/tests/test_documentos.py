import io
import zipfile

import pytest
from docx import Document

from documentos.contextos import bases_legales_unicas, datos_instalacion, resumen_progreso
from documentos.generador import VerticalNoSoportadoError, generar_documento
from documentos.schemas import GenerarDocumentoInput


def _payload(tipo="plan", tipo_instalacion="fotovoltaica_autoconsumo", **extra_exp):
    expediente = {
        "id": "a1b2c3d4-0000-0000-0000-000000000000",
        "tipo_instalacion": tipo_instalacion,
        "comunidad": "andalucia",
        "municipio": "Sevilla",
        "potencia_kw": 8.5,
        "uso": "residencial",
        "referencia_cliente": "Chalet Feria 12",
        **extra_exp,
    }
    plan = {
        "tramites": [
            {
                "orden": 1,
                "nombre": "Memoria Técnica de Diseño",
                "organismo": "Técnico competente",
                "base_legal": "RD 244/2019",
                "plazo_estimado_dias": 7,
                "plazo_legal_dias": None,
                "documentos_requeridos": [
                    {"id": "mtd", "label": "Memoria técnica firmada",
                     "descripcion": "Según modelo autonómico", "obligatorio": True},
                ],
                "notas": None, "plataforma": None, "plataforma_url": None,
                "coste_estimado": "150 EUR", "formulario_ref": "9588",
            },
            {
                "orden": 2,
                "nombre": "Comunicación PUES",
                "organismo": "Junta de Andalucía",
                "base_legal": "Decreto 59/2005",
                "plazo_estimado_dias": 1,
                "plazo_legal_dias": 45,
                "documentos_requeridos": [],
                "notas": None, "plataforma": "PUES", "plataforma_url": None,
                "coste_estimado": None, "formulario_ref": None,
            },
        ],
        "tiempo_total_estimado_dias": 8,
        "advertencias": ["El tiempo total es orientativo y asume trámites en serie."],
    }
    return GenerarDocumentoInput(
        tipo=tipo,
        organizacion={"nombre": "Gestoría Pérez SL", "plan": "pro"},
        expediente=expediente,
        plan=plan,
        tramites_estado={"1": {"estado": "completado", "fecha_inicio": "2026-07-01",
                               "fecha_completado": "2026-07-05"}},
    )


# ── Contextos (lógica pura) ─────────────────────────────────────────────────

def test_datos_instalacion_incluye_referencia_primero():
    filas = datos_instalacion(_payload().expediente)
    assert filas[0] == ("Referencia de cliente", "Chalet Feria 12")
    assert any("8.5 kW" in valor for _, valor in filas)


def test_bases_legales_deduplicadas_y_sin_genericas():
    payload = _payload()
    payload.plan.tramites[1].base_legal = "RD 244/2019"  # duplicada
    bases = bases_legales_unicas(payload)
    assert bases == ["RD 244/2019"]


def test_resumen_progreso():
    assert resumen_progreso(_payload()) == (1, 2)


# ── Render (smoke) ──────────────────────────────────────────────────────────

@pytest.mark.parametrize("tipo,firma", [("plan", b"%PDF"), ("checklist", b"%PDF")])
def test_pdfs_generan_bytes_validos(tipo, firma):
    contenido, media_type, filename = generar_documento(_payload(tipo=tipo))
    assert contenido.startswith(firma)
    assert media_type == "application/pdf"
    assert filename.endswith(".pdf")
    assert len(contenido) > 1000


def test_mtd_docx_contiene_datos_del_expediente():
    contenido, media_type, filename = generar_documento(_payload(tipo="mtd"))
    assert filename.endswith(".docx")
    doc = Document(io.BytesIO(contenido))
    texto = "\n".join(p.text for p in doc.paragraphs)
    celdas = "\n".join(c.text for t in doc.tables for r in t.rows for c in r.cells)
    assert "BORRADOR" in texto
    assert "Sevilla" in celdas
    assert "[POR COMPLETAR]" in celdas


def test_mtd_rechaza_vertical_no_soportado():
    with pytest.raises(VerticalNoSoportadoError):
        generar_documento(_payload(tipo="mtd", tipo_instalacion="gas_baja_presion",
                                   combustible="gas_natural"))


def test_dossier_zip_contenido():
    contenido, media_type, filename = generar_documento(_payload(tipo="dossier"))
    assert media_type == "application/zip"
    with zipfile.ZipFile(io.BytesIO(contenido)) as zf:
        nombres = zf.namelist()
    assert len(nombres) == 3  # plan + checklist + mtd (FV)
    assert any(n.endswith("_MTD_borrador.docx") for n in nombres)


def test_dossier_sin_mtd_para_gas():
    contenido, _, _ = generar_documento(
        _payload(tipo="dossier", tipo_instalacion="gas_baja_presion",
                 combustible="gas_natural")
    )
    with zipfile.ZipFile(io.BytesIO(contenido)) as zf:
        assert len(zf.namelist()) == 2


# ── Integración ligera con el motor: passthrough de formulario_ref ─────────

def test_motor_propaga_formulario_ref():
    from motor_normativo.clasificador import Clasificador
    from schemas.clasificador import ClasificadorInput

    resultado = Clasificador().clasificar(ClasificadorInput(
        tipo_instalacion="fotovoltaica_autoconsumo",
        comunidad="andalucia",
        potencia_kw=600,   # dispara AND-FV-004, cuyos trámites llevan formulario_ref
        uso="industrial",
        municipio="Sevilla",
    ))
    assert any(t.formulario_ref for t in resultado.tramites), (
        "El motor sigue descartando formulario_ref de los JSON de reglas"
    )
