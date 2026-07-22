"""Borrador de Memoria Técnica de Diseño (DOCX) para FV e IRVE.

Genera la estructura y los datos administrativos/técnicos ya conocidos por el
expediente; el contenido técnico (cálculos, esquemas) y la firma corresponden
al técnico competente. El documento lo deja claro en portada.
"""
import io
from datetime import date

from docx import Document
from docx.shared import Pt, RGBColor

from .contextos import (
    VERTICALES_MTD,
    bases_legales_unicas,
    datos_instalacion,
    etiqueta_tipo,
)
from .schemas import GenerarDocumentoInput

POR_COMPLETAR = "[POR COMPLETAR]"

NORMATIVA_BASE = {
    "fotovoltaica_autoconsumo": [
        "RD 842/2002 - Reglamento Electrotécnico para Baja Tensión (REBT)",
        "RD 244/2019 - Condiciones administrativas y técnicas del autoconsumo",
    ],
    "irve": [
        "RD 842/2002 - Reglamento Electrotécnico para Baja Tensión (REBT)",
        "ITC-BT-52 - Instalaciones con fines especiales: infraestructura de recarga de VE",
    ],
}


def _titulo(doc: Document, texto: str, nivel: int = 1):
    doc.add_heading(texto, level=nivel)


def _parrafo_aviso(doc: Document, texto: str):
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xB4, 0x23, 0x18)


def _tabla_pares(doc: Document, filas: list[tuple[str, str]]):
    tabla = doc.add_table(rows=0, cols=2)
    tabla.style = "Table Grid"
    for etiqueta, valor in filas:
        celdas = tabla.add_row().cells
        celdas[0].paragraphs[0].add_run(etiqueta).bold = True
        celdas[1].text = valor


def generar_mtd_docx(payload: GenerarDocumentoInput) -> bytes:
    exp = payload.expediente
    if exp.tipo_instalacion not in VERTICALES_MTD:
        raise ValueError(
            f"El borrador de MTD solo está disponible para fotovoltaica e IRVE "
            f"(recibido: {exp.tipo_instalacion})"
        )

    doc = Document()

    _titulo(doc, "MEMORIA TÉCNICA DE DISEÑO — BORRADOR", 0)
    doc.add_paragraph(etiqueta_tipo(exp.tipo_instalacion))
    _parrafo_aviso(
        doc,
        "BORRADOR generado automáticamente con PermitFlow ES. Requiere la revisión, "
        "cumplimentación de los apartados técnicos y firma de un instalador habilitado "
        "o técnico competente antes de su presentación.",
    )

    _titulo(doc, "1. Datos de la instalación")
    _tabla_pares(doc, datos_instalacion(exp))

    _titulo(doc, "2. Titular de la instalación")
    _tabla_pares(doc, [
        ("Nombre / Razón social", POR_COMPLETAR),
        ("NIF / CIF", POR_COMPLETAR),
        ("Dirección del emplazamiento", f"{POR_COMPLETAR} — {exp.municipio}"),
        ("Teléfono / Email de contacto", POR_COMPLETAR),
    ])

    _titulo(doc, "3. Empresa instaladora habilitada")
    _tabla_pares(doc, [
        ("Razón social", POR_COMPLETAR),
        ("Nº de registro (RII / registro autonómico)", POR_COMPLETAR),
        ("Instalador habilitado (nombre y categoría)", POR_COMPLETAR),
    ])

    _titulo(doc, "4. Normativa de aplicación")
    for norma in NORMATIVA_BASE.get(exp.tipo_instalacion, []):
        doc.add_paragraph(norma, style="List Bullet")
    for base in bases_legales_unicas(payload):
        if base not in NORMATIVA_BASE.get(exp.tipo_instalacion, []):
            doc.add_paragraph(base, style="List Bullet")

    _titulo(doc, "5. Descripción técnica")
    doc.add_paragraph(
        f"{POR_COMPLETAR}: características de los equipos, esquema unifilar, "
        "protecciones, secciones de conductores y justificación de cálculos."
    )
    if exp.tipo_instalacion == "irve":
        doc.add_paragraph(
            f"Esquema de instalación según ITC-BT-52 (colectivo/individual): {POR_COMPLETAR}."
        )
    if exp.tipo_instalacion == "fotovoltaica_autoconsumo":
        doc.add_paragraph(
            f"Modalidad de autoconsumo (con/sin excedentes, RD 244/2019): {POR_COMPLETAR}."
        )

    _titulo(doc, "6. Trámites administrativos asociados")
    filas = [(f"{t.orden}. {t.nombre}", t.organismo) for t in payload.plan.tramites]
    _tabla_pares(doc, filas)

    _titulo(doc, "7. Declaración y firma")
    doc.add_paragraph(
        "El técnico/instalador abajo firmante declara que los datos consignados son "
        "ciertos y que la instalación descrita cumple la normativa de aplicación."
    )
    doc.add_paragraph(f"\nEn {exp.municipio}, a {date.today().strftime('%d/%m/%Y')}")
    doc.add_paragraph("\n\nFdo.: " + POR_COMPLETAR)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
